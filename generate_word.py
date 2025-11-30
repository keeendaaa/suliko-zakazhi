#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для генерации Word документа из JSON меню
"""

import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_menu_document(json_path, output_path):
    """Создает Word документ из JSON меню"""
    
    # Читаем JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Создаем документ
    doc = Document()
    
    # Настройка шрифтов для кириллицы
    def set_font(run, font_name='Times New Roman', size=12):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
    
    # Заголовок ресторана
    title = doc.add_heading(data['restaurant'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, 'Times New Roman', 24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(139, 69, 19)  # Коричневый цвет
    
    doc.add_paragraph()  # Пустая строка
    
    # Обрабатываем каждую категорию
    for category_data in data['menu']:
        # Название категории
        category_heading = doc.add_heading(category_data['category'], level=1)
        for run in category_heading.runs:
            set_font(run, 'Times New Roman', 18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(178, 34, 34)  # Красный цвет
        
        # Описание категории (если есть)
        if category_data.get('description'):
            desc_para = doc.add_paragraph(category_data['description'])
            desc_para.style = 'Normal'
            for run in desc_para.runs:
                set_font(run, 'Times New Roman', 11)
                run.font.italic = True
            doc.add_paragraph()  # Пустая строка
        
        # Блюда в категории
        for item in category_data['items']:
            # Название блюда
            item_para = doc.add_paragraph()
            item_run = item_para.add_run(item['name'])
            set_font(item_run, 'Times New Roman', 14)
            item_run.font.bold = True
            
            # Описание (если есть)
            if item.get('description'):
                desc_run = item_para.add_run(f" — {item['description']}")
                set_font(desc_run, 'Times New Roman', 12)
                desc_run.font.italic = True
            
            # Порция и цена
            price_para = doc.add_paragraph()
            price_text = f"{item['portion']} | {item['price']}₽"
            price_run = price_para.add_run(price_text)
            set_font(price_run, 'Times New Roman', 12)
            price_run.font.bold = True
            price_run.font.color.rgb = RGBColor(139, 69, 19)  # Коричневый цвет
            
            doc.add_paragraph()  # Пустая строка между блюдами
        
        # Заметки (если есть)
        if category_data.get('notes'):
            doc.add_paragraph()
            notes_para = doc.add_paragraph('Примечания:')
            for run in notes_para.runs:
                set_font(run, 'Times New Roman', 12)
                run.font.bold = True
            
            for key, value in category_data['notes'].items():
                note_para = doc.add_paragraph(f"• {value}", style='List Bullet')
                for run in note_para.runs:
                    set_font(run, 'Times New Roman', 11)
                    run.font.italic = True
        
        # Разделитель между категориями
        doc.add_paragraph()
        doc.add_paragraph('─' * 50)
        doc.add_paragraph()
    
    # Цитаты
    if data.get('quotes'):
        doc.add_page_break()
        quotes_heading = doc.add_heading('Цитаты', level=1)
        for run in quotes_heading.runs:
            set_font(run, 'Times New Roman', 18)
            run.font.bold = True
        
        for quote in data['quotes']:
            quote_para = doc.add_paragraph(quote['text'])
            quote_para.style = 'Quote'
            for run in quote_para.runs:
                set_font(run, 'Times New Roman', 12)
                run.font.italic = True
                run.font.color.rgb = RGBColor(139, 69, 19)
            
            if quote.get('category'):
                cat_para = doc.add_paragraph(f"— {quote['category']}")
                for run in cat_para.runs:
                    set_font(run, 'Times New Roman', 10)
                    run.font.italic = True
            
            doc.add_paragraph()
    
    # Сохраняем документ
    doc.save(output_path)
    print(f"Word документ создан: {output_path}")

if __name__ == '__main__':
    import sys
    json_path = 'menu.json'
    output_path = 'menu.docx'
    
    if len(sys.argv) > 1:
        json_path = sys.argv[1]
    if len(sys.argv) > 2:
        output_path = sys.argv[2]
    
    create_menu_document(json_path, output_path)

